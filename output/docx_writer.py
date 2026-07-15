import os
import re
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def write_docx(notes: dict, frames: list[dict], output_dir: str, safe_title: str,
               log_cb=None, source_info: dict | None = None) -> str:
    """Build the DOCX from structured notes and return the output file path."""
    doc = Document()
    _setup_styles(doc)

    _add_title(doc, notes.get("title", safe_title))
    _add_source_block(doc, source_info)

    frames_by_idx = {i + 1: f["path"] for i, f in enumerate(frames)}
    boxes = notes.get("screenshot_boxes") or {}

    for chapter in notes.get("chapters", []):
        _add_chapter(doc, chapter, frames_by_idx, boxes, log_cb)

    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{safe_title}_notes.docx")
    doc.save(out_path)
    return out_path


def _setup_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_title(doc: Document, title: str):
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    doc.add_paragraph()


def _add_source_block(doc: Document, source_info: dict | None):
    """Recording metadata at the top of the document: type, URL/path, and
    the recording summary (YouTube description / Teams AI recap) if any."""
    if not source_info:
        return

    rec_type = (source_info.get("type") or "").strip()
    url = (source_info.get("url") or "").strip()
    summary = (source_info.get("summary") or "").strip()

    if rec_type:
        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True
        p.add_run(rec_type)

    if url:
        p = doc.add_paragraph()
        p.add_run("Recording: ").bold = True
        if url.startswith(("http://", "https://")):
            _add_hyperlink(p, url, url)
        else:
            p.add_run(url)

    if summary:
        p = doc.add_paragraph()
        p.add_run("Recording summary:").bold = True
        blank = False
        for line in summary.splitlines():
            line = line.rstrip()
            if not line:
                blank = True   # collapse runs of blank lines to one break
                continue
            sp = doc.add_paragraph()
            if blank:
                sp.paragraph_format.space_before = Pt(6)
                blank = False
            run = sp.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable external hyperlink (python-docx has no API for this)."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)


def _add_chapter(doc: Document, chapter: dict, frames_by_idx: dict,
                 boxes: dict, log_cb=None):
    heading = doc.add_heading(chapter.get("title", ""), level=1)
    heading.runs[0].font.size = Pt(16)

    speakers = chapter.get("speakers", [])
    if speakers:
        p = doc.add_paragraph()
        run = p.add_run(f"Speakers: {', '.join(speakers)}")
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(10)

    for point in chapter.get("key_points", []):
        screenshot_idx = point.get("screenshot_idx")

        if screenshot_idx is not None and screenshot_idx in frames_by_idx:
            _add_screenshot(doc, frames_by_idx[screenshot_idx],
                            boxes.get(screenshot_idx), log_cb)

        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        _add_formatted_run(p, point.get("text", ""))

    doc.add_paragraph()


def _add_screenshot(doc: Document, image_path: str, box=None, log_cb=None):
    """Embed a screenshot, cropped to its content_box when one was provided.

    Fallback chain: cropped → full-frame Pillow re-encode (also covers JPEGs
    without a JFIF marker that python-docx rejects) → skip with a log line.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    try:
        if box:
            run.add_picture(_image_stream(image_path, box), width=Inches(5.5))
        else:
            run.add_picture(image_path, width=Inches(5.5))
    except Exception as e:
        try:
            run.add_picture(_image_stream(image_path, None), width=Inches(5.5))
        except Exception:
            if log_cb:
                log_cb(f"Skipped screenshot {image_path}: {type(e).__name__}: {e}")


def _image_stream(image_path: str, box):
    """Return a JPEG BytesIO of the image, cropped to the fractional
    [x0, y0, x1, y1] box when given."""
    import io
    from PIL import Image

    img = Image.open(image_path)
    if box:
        w, h = img.size
        x0, y0, x1, y1 = box
        img = img.crop((round(x0 * w), round(y0 * h), round(x1 * w), round(y1 * h)))
    buf = io.BytesIO()
    img.convert("RGB").save(buf, "JPEG", quality=92)
    buf.seek(0)
    return buf


def _add_formatted_run(paragraph, text: str):
    """Parse **bold** markers and add appropriately styled runs."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
