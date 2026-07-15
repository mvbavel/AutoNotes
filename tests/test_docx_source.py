"""Tests for the recording-source block at the top of the DOCX."""
import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from output.docx_writer import write_docx

NOTES = {"title": "T", "chapters": [{"title": "C", "start_time": 0, "end_time": 1,
                                     "speakers": [], "key_points": [
                                         {"text": "point", "screenshot_idx": None}]}]}


def _build(source_info):
    tmp = tempfile.mkdtemp()
    path = write_docx(dict(NOTES), [], tmp, "t", source_info=source_info)
    return Document(path)


class TestSourceBlock(unittest.TestCase):
    def test_full_block(self):
        doc = _build({"type": "YouTube",
                      "url": "https://youtu.be/abc123",
                      "summary": "Line one.\n\n\nLine two."})
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Source: YouTube", text)
        self.assertIn("Recording:", text)
        self.assertIn("Line one.", text)
        self.assertIn("Line two.", text)
        # URL is a real external hyperlink relationship
        rels = [r for r in doc.part.rels.values() if r.reltype == RT.HYPERLINK]
        self.assertEqual(len(rels), 1)
        self.assertEqual(rels[0].target_ref, "https://youtu.be/abc123")

    def test_local_file_path_is_plain_text(self):
        doc = _build({"type": "Local file",
                      "url": "/Users/x/Movies/demo.mp4", "summary": ""})
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Source: Local file", text)
        self.assertIn("/Users/x/Movies/demo.mp4", text)
        self.assertIn("point", text)
        rels = [r for r in doc.part.rels.values() if r.reltype == RT.HYPERLINK]
        self.assertEqual(rels, [])
        self.assertNotIn("Recording summary", text)

    def test_no_source_info_is_clean(self):
        doc = _build(None)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("Source:", text)
        self.assertIn("point", text)


if __name__ == "__main__":
    unittest.main()
